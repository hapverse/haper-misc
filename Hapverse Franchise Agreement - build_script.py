# -*- coding: utf-8 -*-
"""Builds the redrafted Hapverse Franchise Agreement as a formatted .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- base styles ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def P(text, bold=False, italic=False, align=None, size=None, space_after=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size: r.font.size = Pt(size)
    if align == 'c': p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'l': p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def CL(num, text):
    """Sub-clause: bold number then body."""
    p = doc.add_paragraph()
    rn = p.add_run(num + '  ')
    rn.bold = True
    p.add_run(text)
    return p

def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def TITLE(text, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    return p

def HR():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '999999')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

# ---------- footer with page number ----------
def _add_field(paragraph, field):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = field
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(e)
    return run

_footer = doc.sections[0].footer
_fp = _footer.paragraphs[0]
_fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_r = _fp.add_run('HAPVERSE PRIVATE LIMITED — Franchise Agreement   |   Page ')
_r.font.size = Pt(8); _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
_pg = _add_field(_fp, 'PAGE'); _pg.font.size = Pt(8); _pg.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
_r2 = _fp.add_run(' of '); _r2.font.size = Pt(8); _r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
_tp = _add_field(_fp, 'NUMPAGES'); _tp.font.size = Pt(8); _tp.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# =====================================================================
# COVER
# =====================================================================
TITLE('FRANCHISE AGREEMENT', 18)
TITLE('HAPER QUICK-COMMERCE PLATFORM', 12)
P('HAPVERSE PRIVATE LIMITED', bold=True, align='c')
HR()
P('Date of Execution: ____________________        Effective Date: ____________________', align='c')
P('Place of Execution: Bengaluru, Karnataka', align='c')
P('')

# Drafting / use note
note = doc.add_paragraph()
note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
nr = note.add_run('NOTE FOR THE COMPANY (delete before issuing to a franchisee): Blanks shown as "____" or "[ ]" are to be filled in before signing. Complete the "Schedule of Items to Fill Before Signing" at the end. This Agreement and the separate Deed of Personal Guarantee (Annexure C) must each be stamped under the Stamp Act of the State where it is executed (Karnataka if signed at Bengaluru; Bihar if signed in Bihar) and signed by two witnesses before they are relied upon. See also the GST-invoicing decision in Clause 6 and the foreign-investment note in the "Schedule of Items to Fill Before Signing".')
nr.italic = True
nr.font.size = Pt(8.5)
nr.font.color.rgb = RGBColor(0x88, 0x44, 0x00)

# =====================================================================
# PARTIES
# =====================================================================
H1('PARTIES')
P('This FRANCHISE AGREEMENT (the "Agreement") is made and entered into at Bengaluru, Karnataka on the Execution Date stated above.')
P('BY AND BETWEEN:', bold=True)
CL('(1)', 'HAPVERSE PRIVATE LIMITED, a company incorporated under the Companies Act, 2013 on 16 December 2025 (CIN: U46901KA2025PTC212605; PAN: AAICH3580R; TAN: BLRH18936B; GSTIN: 29AAICH3580R1ZP), having its registered office at No. 472/7 Balaji Arcade, 2nd & 3rd Floor, A.V.S. Compound, 20th L Cross Road, AVS Layout, Ejipura, Koramangala 4th Block, Bengaluru, Karnataka – 560095, represented by its authorised signatory (the "Company" or "Hapverse", which expression includes its successors and permitted assigns) of the FIRST PART;')
P('AND', bold=True)
CL('(2)', 'Mr./Ms. ____________________________ , son/daughter/wife of ____________________ , aged ____ years, an individual [carrying on business as a sole proprietor under the name "____________________"], holding PAN ____________ and Aadhaar No. XXXX XXXX ________ (last 4 digits shown; the Company retains only a masked/redacted copy with consent, in line with UIDAI norms, and not the full Aadhaar number), GSTIN (if registered): ____________ , residing at ____________________________________ (the "Franchisee", which expression includes his/her heirs, executors, administrators, legal representatives and permitted assigns) of the SECOND PART.')
P('The Company and the Franchisee are referred to collectively as the "Parties" and individually as a "Party".')

H2('RECITALS')
CL('A.', 'The Company owns and operates the "Haper" quick-commerce and grocery-delivery platform, comprising the Haper mobile application, an order and inventory management system, an admin dashboard, related technology, and the "HAPER" and "HAPVERSE" brands (together, the "Haper System").')
CL('B.', 'The Company appoints franchisees to operate "Haper"-branded stores from which orders placed on the platform are fulfilled, on the terms of this Agreement.')
CL('C.', 'The Franchisee has represented that it has the resources, a suitable premises recommendation, and the capability to operate such a store, and wishes to be appointed on these terms. Relying on those representations, the Company is willing to grant the franchise.')
P('NOW THEREFORE, in consideration of the mutual covenants below, the Parties agree as follows:', bold=True)

# =====================================================================
# 1. DEFINITIONS
# =====================================================================
H1('1. DEFINITIONS AND INTERPRETATION')
defs = [
 ('1.1', '"Applicable Law" means all Indian statutes, rules, regulations, notifications, guidelines and orders of any government or regulatory authority, as amended from time to time.'),
 ('1.2', '"Business" means the operation of a "Haper"-branded quick-commerce store and the fulfilment of grocery and related orders placed through the Haper System.'),
 ('1.3', '"Confidential Information" means all non-public information disclosed by or on behalf of the Company, including vendor data, pricing, business and expansion plans, Customer Data, the SOPs, software, technology, processes, and marketing strategies, in any form, and whether or not marked confidential.'),
 ('1.4', '"Customer Data" means all personal and transactional data of customers generated, collected or processed through the Haper System or at the Store.'),
 ('1.5', '"Effective Date" means the date stated on the cover page or, if that date is left blank, the date on which the first installment of the Security Deposit is received in cleared funds by the Company.'),
 ('1.6', '"Events of Default" has the meaning given in Clause 18.2.'),
 ('1.7', '"Security Deposit" means the refundable, interest-free deposit of ₹7,50,000 payable under Clause 4.1.'),
 ('1.8', '"GMV" (Gross Merchandise Value) means the total value of successfully completed customer orders fulfilled from the Store, whether placed through the Haper System or otherwise and whether paid digitally or in cash, excluding only cancelled orders, bona fide refunds, fraudulent orders and chargebacks.'),
 ('1.9', '"Initial Term" means the period stated in Clause 3.1.'),
 ('1.10', '"Licensed Marks" means the trade marks, names and logos "HAPER" and "HAPVERSE" and any other marks the Company notifies, particulars of which are set out in (or to be inserted in) Annexure B.'),
 ('1.11', '"SOP" means the Standard Operating Procedures issued and updated by the Company (summarised in Annexure D).'),
 ('1.12', '"Store" means the premises approved by the Company and described in Annexure B from which the Business is conducted.'),
 ('1.13', '"Territory" means the geographical area within a radius of 2 (two) kilometres from the Store, or such other radius (not exceeding 3 (three) kilometres) as is specified in Annexure B.'),
]
for n, t in defs: CL(n, t)
CL('1.14', 'Interpretation: clause headings are for convenience only; the singular includes the plural; "including" means "including without limitation"; references to a statute include subordinate legislation and amendments; and in case of conflict the body of this Agreement prevails over the Annexures, save for the specific commercial particulars set out in Annexure A (revenue) and Annexure B (territory and Store).')

# =====================================================================
# 2. GRANT & TERRITORY
# =====================================================================
H1('2. GRANT OF FRANCHISE AND TERRITORY')
CL('2.1', 'The Company grants the Franchisee a non-exclusive (save as provided in Clause 2.2), non-transferable (except as permitted under Clause 21) right to operate one (1) "Haper"-branded Store within the Territory during the Term, using the Haper System in accordance with this Agreement and the SOPs.')
CL('2.2', 'Conditional Territorial Exclusivity: So long as the Franchisee is not in material breach and has not lost exclusivity under Clause 11.3, the Company will not itself operate, or appoint another franchisee to operate, a substantially similar "Haper"-branded Store within the Territory. This exclusivity is a privilege conditional on performance and converts to non-exclusive on sustained shortfall as provided in Clause 11.3, without affecting any other remedy.')
CL('2.3', 'The Franchisee shall not operate, solicit or fulfil orders outside the Territory without the Company’s prior written approval.')
CL('2.4', 'The exclusivity in Clause 2.2 does not restrict the Company from serving customers anywhere (including within the Territory) through its own corporate, institutional, marketplace, dark-store, or direct online channels, and no such activity is a breach of this Agreement.')
CL('2.5', 'The Company may establish, operate or franchise any number of stores outside the Territory.')

# =====================================================================
# 3. TERM & RENEWAL
# =====================================================================
H1('3. TERM AND RENEWAL')
CL('3.1', 'Initial Term: This Agreement is valid for 3 (three) years from the Effective Date, unless terminated earlier in accordance with its terms (the "Initial Term"; together with any renewal, the "Term").')
CL('3.2', 'Renewal: Renewal is not automatic. Any renewal is subject to mutual written agreement, execution of a fresh agreement, a satisfactory compliance and performance review, payment of a renewal fee (if any), and such revised commercial terms as the Company may reasonably specify.')
CL('3.3', 'Expansion: Any expansion of inventory capacity, infrastructure, delivery fleet or technology footprint may require additional investment by the Franchisee, subject to mutual written agreement.')

# =====================================================================
# 4. FEE & FINANCIAL TERMS
# =====================================================================
H1('4. SECURITY DEPOSIT AND FINANCIAL TERMS')
CL('4.1', 'Security Deposit: The Franchisee shall pay a refundable, interest-free Security Deposit of ₹7,50,000 (Rupees Seven Lakh Fifty Thousand Only) as security for the due performance of its obligations, including its inventory and asset liabilities under Clauses 6 and 10, payable in 3 (three) equal installments of ₹2,50,000 (Rupees Two Lakh Fifty Thousand Only) each: (i) the first on or before the Effective Date; (ii) the second on or before the 15th (fifteenth) day from the Effective Date; and (iii) the third on or before the 30th (thirtieth) day from the Effective Date. The Security Deposit is not a fee and is not consideration for any supply; accordingly, no GST is chargeable on it when taken or when refunded (Section 2(31), CGST Act, 2017). It is refundable in accordance with Clause 19, after the deductions stated there. No interest is payable on the Security Deposit.')
CL('4.2', 'Payment and go-live gate: Each installment is payable by electronic bank transfer (no cash) to — Account Name: HAPVERSE PRIVATE LIMITED; Account No.: 4052802360 (Current Account); Bank & Branch: Kotak Mahindra Bank, Bengaluru – Koramangala; IFSC: KKBK0000424 — against a receipt (and not a tax invoice) issued for each installment. Notwithstanding anything else, the Company shall not stock inventory at, or commence live operations (go-live) of, the Store until the full Security Deposit (all three installments) has been received in cleared funds. If any installment is not paid by its due date, the Company may suspend setup, withhold go-live and/or terminate this Agreement, and shall refund the amounts already received less its actual documented costs incurred and any dues. The Company is under no obligation to execute the Store lease or commence fit-out until the full Security Deposit is received; if it elects to do so earlier and any installment is then not paid, the Franchisee is liable for the Company’s full resulting exposure (including committed lease and fit-out costs), recoverable under Clause 4.5.')
CL('4.3', 'What the Company provides: In consideration of the mutual covenants in this Agreement (and not of the Security Deposit), the Company grants the franchise and provides technology access, brand-usage rights, training, launch support and initial store-setup and operational guidance, and leases and equips the Store premises under Clause 10. The Franchisee shall identify and recommend suitable Store premises for the Company’s approval. The Company’s remuneration is the revenue share under Clause 11 and Annexure A.')
CL('4.4', 'GST and Taxes: Amounts payable for taxable supplies — such as the Franchisee’s commission/support service, and any amount later adjusted from the Security Deposit against a taxable supply — are exclusive of GST, which the payer shall bear additionally at the applicable rate against a valid tax invoice. Tax shall be deducted at source where required at the rates in force from time to time.')
CL('4.5', 'Additional Security and KYC: Before launch, the Franchisee shall deliver to the Company (i) self-attested KYC — PAN, a recent photograph, an independent address proof, a cancelled cheque of the account to which the Security Deposit is to be refunded, and a masked Aadhaar (last 4 digits) or other government photo-ID; and (ii) 3 (three) signed security cheques and a signed e-NACH/auto-debit mandate on that account, to secure payment of any dues, deductions or losses exceeding the Security Deposit. The Franchisee authorises the Company, under Section 20 of the Negotiable Instruments Act, 1881, to fill in the date and the amount (up to the finally-determined dues) on such cheques and to present them with a statement of dues, the Franchisee acknowledging that each cheque is given towards an existing and continuing liability. The Company may present such cheques or invoke the mandate to recover any admitted or finally-determined dues. The Franchisee, being an individual / sole proprietor, is personally liable without limit for all obligations under this Agreement, and the Company may proceed against the Franchisee’s personal assets for any amount exceeding the Security Deposit.')

# =====================================================================
# 5. LOCK-IN
# =====================================================================
H1('5. LOCK-IN PERIOD')
CL('5.1', 'The Parties agree to a lock-in period of 12 (twelve) months from the Effective Date.')
CL('5.2', 'During the lock-in period the Franchisee shall not voluntarily terminate this Agreement without the Company’s written consent, and may not serve a notice of termination for convenience that would take effect during the lock-in period. The Company’s right to terminate for cause is not restricted by the lock-in period.')

# =====================================================================
# 6. INVENTORY + GST CHARACTERISATION
# =====================================================================
H1('6. INVENTORY OWNERSHIP, GST CHARACTERISATION AND MANAGEMENT')
CL('6.1', 'Ownership: All inventory supplied to the Store remains the sole and absolute property of the Company until it is sold to a customer. Title never passes to the Franchisee. The Franchisee holds such inventory as the Company’s bailee and custodian within the meaning of Sections 148 to 152 of the Indian Contract Act, 1872, and not as owner.')
CL('6.2', 'Custodian and Commission-Agent Role: The Franchisee operates the Store and acts as the Company’s commission agent and custodian. At no time does the Franchisee purchase inventory for resale on its own account.')
CL('6.3', 'GST Characterisation and Consumer Invoicing: As owner of the inventory, the Company alone determines and may vary retail prices, discounts and offers; the Franchisee shall not independently set, alter or discount prices, nor sell otherwise than through the Haper System. ALL tax invoices to end-customers shall be raised in the name and under the GSTIN of the Company, the Store being registered as the Company’s additional place of business (or, where in a different State, under a separate registration). The Franchisee shall not raise any sale invoice to customers in its own name. The Company is solely liable for output GST on customer supplies. The Franchisee’s only outward supply to the Company is its commission/support service, for which it shall raise a valid GST tax invoice (a precondition to payout and to the Company’s input tax credit). The Franchisee shall maintain records and e-way-bill compliance as the Company directs, and shall indemnify the Company for any additional GST, interest or penalty arising from the Franchisee invoicing in its own name contrary to this Clause. Any invoicing by the Franchisee to customers in its own name is a material breach and an immediate-termination event under Clause 18.1.')
CL('6.4', 'Inventory Losses: The Franchisee is liable for the Company’s actual documented cost (landed cost net of any return credit) of stock that is damaged, missing, stolen, spoiled or lost while in its custody due to its act, omission, negligence or mishandling, plus any input tax credit the Company is required to reverse on such stock under Section 17(5)(h) of the CGST Act, 2017. Such liability is an indemnity for the Company’s actual, documented loss.')
CL('6.5', 'Expiry Management: The Franchisee shall monitor expiry dates and return or exchange near-expiry products as per the SOPs. If, in breach of the SOPs, the Franchisee fails to do so, it shall bear the Company’s actual documented landed cost of the expired stock.')
CL('6.6', 'All-Sales-Through-Platform: The Franchisee shall route every sale from the Store through the Haper System, shall stock and sell only Company-approved goods, and shall not sell competing or its own goods from the Store.')
CL('6.7', 'Cash (COD) Handling: All cash collected on delivery is the Company’s money, held by the Franchisee in trust, kept segregated, reconciled daily on the Admin Dashboard, and remitted to the Company by the next business day. Under-reporting, diversion or manipulation of sales, cash or inventory is treated as inventory manipulation and fraud under Clause 18.1. COD and other trust monies are the Company’s property, are not subject to any set-off or retention by the Franchisee, and may be appropriated by the Company against any dues. Any unexplained shortfall or variance in cash, sales or stock beyond the tolerance specified in the SOPs is rebuttably presumed attributable to the Franchisee; the Franchisee shall furnish photographic/video evidence and obtain the Company’s written sign-off before any expiry, spoilage or shrinkage credit is allowed.')

# =====================================================================
# 7. COMPANY RESPONSIBILITIES
# =====================================================================
H1('7. RESPONSIBILITIES OF THE COMPANY')
P('The Company shall provide:')
CL('7.1', 'Technology: access to the Haper mobile app, Admin Dashboard and order/inventory management systems, with reasonable technical support.')
CL('7.2', 'Marketing: centralised brand assets, digital-marketing support and promotional campaigns.')
CL('7.3', 'Operations: initial onboarding, staff-training modules, SOP documentation and delivery-management support.')
CL('7.4', 'Store Launch and Initial 30-Day Support: The Company shall make the Store operational (go-live) from the first day of its launch, and shall provide dedicated launch and operational support for a period of 30 (thirty) days from the launch date — including activation of the systems and the Haper app, initial inventory stocking, staff hand-holding, and resolution of teething operational issues — so that the Store can run smoothly. Thereafter the Company shall continue to provide ongoing technology, marketing and operational support under Clauses 7.1 to 7.3. Launch is subject to receipt of the full Security Deposit in accordance with Clause 4.2.')
P('The Company does not warrant uninterrupted or error-free operation of the platform.')

# =====================================================================
# 8. FRANCHISEE RESPONSIBILITIES
# =====================================================================
H1('8. RESPONSIBILITIES OF THE FRANCHISEE')
P('The Franchisee shall:')
CL('8.1', 'maintain the Store in accordance with brand, hygiene and safety standards and the SOPs;')
CL('8.2', 'hire, manage, pay, supervise and discipline its own staff, and comply with all labour laws (including the Code on Wages, ESI, EPF, Shops & Establishments, Professional Tax and the Sexual Harassment of Women at Workplace (Prevention, Prohibition and Redressal) Act, 2013) as applicable. The Franchisee’s staff are its employees alone and not the Company’s;')
CL('8.3', 'strictly adhere to the SOPs and delivery standards (Annexure D);')
CL('8.4', 'obtain and keep valid at all times every licence and registration required for the Store, including FSSAI licence/registration, GST registration, Shops & Establishments registration, Professional Tax registration, fire and municipal/trade permits, and EPF/ESI registration as applicable, and upload proof of each to the Admin Dashboard before launch and within 7 (seven) days of each renewal;')
CL('8.5', 'process customer personal data only as the Company’s data processor under Clause 13 and Annexure E;')
CL('8.6', 'permit the Company to inspect and audit the Store under Clause 11.6; and')
CL('8.7', 'ensure that every delivery person engaged for the Store holds a valid driving licence and valid vehicle registration and insurance, observes all traffic and road-safety laws, and that no delivery timeline in the SOPs is treated as authorising any unsafe or unlawful act.')

# =====================================================================
# 9. LICENCES & FOOD SAFETY
# =====================================================================
H1('9. LICENCES AND FOOD SAFETY (FSSAI) COMPLIANCE')
CL('9.1', 'The Store is a "food business" under the Food Safety and Standards Act, 2006. As a condition precedent to launch and continuously throughout the Term, the Franchisee shall obtain and maintain a valid FSSAI licence/registration of the appropriate category, display it at the Store and on the platform listing, deliver a certified copy to the Company before launch and within 7 (seven) days of each renewal, and renew it at least 30 (thirty) days before expiry.')
CL('9.2', 'The Franchisee shall comply with the FSS Act and its regulations, including hygiene, cold-chain and food-handler medical-fitness requirements, maintain a FoSTaC-certified Food Safety Supervisor where required, display the Food Safety Display Board, follow a documented FIFO and near-expiry quarantine process, and shall not sell, dispatch or fulfil any expired, recalled, adulterated, mislabelled or otherwise unsafe food.')
CL('9.3', 'Any expiry, suspension, cancellation, prosecution or improvement notice relating to any licence shall be reported to the Company within 24 (twenty-four) hours, and entitles the Company to suspend the Store on the platform and to terminate this Agreement immediately notwithstanding the lock-in period. The Franchisee shall bear recall costs and, where it has breached the SOPs, the Company’s actual landed cost of affected stock.')
CL('9.4', 'Legal Metrology and Consumer Protection: The Franchisee shall handle only correctly labelled goods, shall never charge above the printed MRP, and shall enter accurate product, weight, price and country-of-origin information on the platform. The Franchisee shall comply with the Legal Metrology Act, 2009 and the Consumer Protection (E-Commerce) Rules, 2020 as applicable, and shall handle plastic-packaging/EPR obligations for its operations. The Franchisee indemnifies the Company for any consumer or regulatory action arising from its breach of this Clause.')
CL('9.5', 'The Company shall obtain and maintain its own FSSAI registrations as required of an e-commerce food-business operator (including any central e-commerce FBO licence and storage/warehouse licence for the premises), and shall comply with applicable FSSAI directions (including minimum-shelf-life-on-delivery norms) and the seller-side disclosures required of an inventory e-commerce entity under the Consumer Protection (E-Commerce) Rules, 2020. Nothing in this Agreement transfers or discharges these non-delegable obligations of the Company; the allocation of responsibility operates only as between the Parties and not against any authority or consumer.')

# =====================================================================
# 10. STORE ASSETS
# =====================================================================
H1('10. STORE PREMISES, ASSETS AND INFRASTRUCTURE')
CL('10.1', 'All assets funded or installed by the Company (including racks, refrigerators, POS devices and signage) remain the property of the Company.')
CL('10.2', 'On termination, the Company may recover such assets or deduct their depreciated value from the final settlement. The Franchisee shall insure the Company-owned inventory and assets at the Store against fire, theft and damage, naming the Company as loss-payee, and shall produce proof of cover on request.')
CL('10.3', 'Premises and Lease: The Franchisee shall identify and recommend a suitable premises for the Store for the Company’s prior written approval. The Company shall take the approved premises on a REGISTERED lease for a minimum term of 3 (three) years, in the Company’s name, and shall bear the lease stamp duty, registration charges, the landlord’s security deposit and the monthly rent. The Franchisee shall assist in securing the premises and shall obtain the landlord’s written no-objection for commercial / quick-commerce use, food business (FSSAI licensing), and Haper branding and signage, together with the landlord’s consent to the Company’s occupation and to assignment or sub-licensing within the Haper network. The Company shall, so far as commercially practicable, ensure the lease contains a tenant-side break/early-exit right and the right to assign or sub-licence within the Haper network; the landlord’s security deposit shall be paid by, and refundable directly to, the Company. The Company bears the rent directly to the landlord as part of the support and infrastructure it provides; the Franchisee bears no rent and is paid no separate rent component, and the Support Amount in Annexure A already accounts for the Company bearing the rent and shall not be supplemented by any further rent payment to the Franchisee. If the landlord is unregistered for GST, the Company shall self-invoice and pay GST on the rent under reverse charge (Notification 09/2024-CT(Rate), w.e.f. 10 October 2024) and claim input credit. The Franchisee warrants that it has no direct or indirect interest in, or relationship with, the landlord and that the rent is arm’s-length, indemnifies the Company for breach of this warranty, and shall not, during the Term or for 24 (twenty-four) months after exit, directly or indirectly take, solicit or negotiate any lease or occupancy of the Store premises or deal with the landlord to the Company’s detriment.')
CL('10.4', 'Premises on exit: The Store premises and the lease are, and remain, the Company’s. The Franchisee occupies the Store solely as the Company’s bare licensee / permissive occupant and acquires no tenancy, sub-tenancy or interest in the premises. On termination or closure the Franchisee shall have no right, interest or claim in the premises or the lease, shall hand over vacant operational control of the Store, and shall sign any handover, surrender or assignment documents the Company reasonably requires; the Company may, at its sole discretion, continue to operate the Store, appoint a replacement franchisee at the same premises, or relocate. Failure to hand over vacant control entitles the Company to a mandatory injunction and to holdover charges of twice the per-day rent until vacated, recoverable from the Security Deposit and the Franchisee.')

# =====================================================================
# 11. REVENUE SHARING, PERFORMANCE & AUDIT
# =====================================================================
H1('11. REVENUE SHARING, PAYOUT, PERFORMANCE AND AUDIT')
CL('11.1', 'Minimum Performance Standard: The Store shall achieve a minimum monthly GMV of ₹3,00,000 (Rupees Three Lakh Only) (the "Minimum Performance Standard").')
CL('11.2', 'Payout: Monthly Payout = Sales Commission + Support Amount − permitted Deductions. Commission and Support are as per Annexure A. Payouts are processed on or before the 5th day of the following calendar month, against a valid GST tax invoice from the Franchisee. Where the Franchisee is not required to be, and is not, GST-registered, it shall instead issue a bill of supply or commercial invoice, and the Company shall account for any tax under reverse charge where applicable; payout shall not be withheld solely for the absence of a tax invoice where GST registration is not legally required. The Franchisee shall disclose if it is registered as a micro or small enterprise; the undisputed portion of any payout shall be paid within 45 (forty-five) days regardless of any dispute over deductions, and disputed amounts shall be ring-fenced so as not to delay the undisputed payout.')
CL('11.3', 'Consequences of Shortfall: If the Store fails to meet the Minimum Performance Standard for any 3 (three) consecutive months, then, in addition to any other remedy: (a) the territorial exclusivity in Clause 2.2 automatically converts to non-exclusive; (b) the Company may, on 30 (thirty) days’ written notice, reduce or suspend the Support Amount; and (c) if the shortfall continues for a further 3 (three) consecutive months, the Company may terminate this Agreement for cause under Clause 18, notwithstanding the lock-in period.')
CL('11.4', 'Minimum Operating Covenant: The Franchisee shall keep the Store open and operational during the hours specified in the SOPs, shall maintain the minimum stock/par levels the Company specifies, and shall place orders so as to keep the Store adequately stocked. Failure to keep the Store operational and stocked is a material breach.')
CL('11.5', 'Permitted Deductions: The Company may set off against the Monthly Payout amounts for (i) its documented cost of damaged, missing or expired inventory attributable to the Franchisee; (ii) customer compensation or chargebacks whose root cause is attributable to the Franchisee; and (iii) pre-agreed service-level charges set out in Annexure D, which the Parties agree are genuine pre-estimates of loss and not penalties. The Company shall furnish an itemised deduction statement with each payout; the Franchisee may dispute any deduction within 15 (fifteen) days, and a disputed amount is not final until resolved under Clause 24. No deduction shall be made for matters attributable to the Company’s platform, app, marketing or delivery-technology obligations.')
CL('11.6', 'Audit: The Company (or its representatives) shall conduct a periodic audit of the Store at least once every 6 (six) months (half-yearly), and may conduct it quarterly or more frequently at its discretion — covering a physical stock/inventory count, cash and COD reconciliation, sales records, validity of licences, and operational/SOP compliance — and may, in addition, audit at any time on reasonable notice, or without prior notice where fraud, diversion or under-reporting is suspected. The Franchisee shall cooperate fully and grant access to the premises, stock, POS, bank statements, CCTV and all records, and shall countersign each audit and stock-reconciliation report. Any stock or cash shortfall identified is recoverable as a permitted Deduction under Clause 11.5 and/or from the Security Deposit. If any audit reveals a shortfall, diversion or under-reporting exceeding 2% (two per cent) of audited GMV or audited stock value, the Franchisee shall bear the reasonable cost of that audit, without prejudice to Clauses 18.1 and 19.')

# =====================================================================
# 12. IP & TECH LICENCE
# =====================================================================
H1('12. INTELLECTUAL PROPERTY AND TECHNOLOGY LICENCE')
CL('12.1', 'The Company is the proprietor of, or applicant/authorised licensor for, the Licensed Marks (registration particulars to be inserted in Annexure B) and the Haper System. The Company grants the Franchisee a limited, revocable, non-transferable, non-sub-licensable licence to use the Licensed Marks and the Haper System solely to operate the Business during the Term, strictly in accordance with the SOPs and brand standards.')
CL('12.2', 'Quality Control: The licence is conditional on the Franchisee’s continuing conformity with the SOPs and brand/quality standards and the Company’s audits. All use of the Licensed Marks, and all goodwill arising from such use, enures solely to the benefit of the Company; the Franchisee acquires no right, title, goodwill or interest in them.')
CL('12.3', 'The Franchisee shall not copy, modify, reverse-engineer, scrape, resell or create derivative works of the Haper System or software, nor register or use any confusingly similar mark, domain, app name or social-media handle, and shall not challenge the validity of or the Company’s title to the Licensed Marks or the Haper System.')
CL('12.4', 'The Company’s software, databases and the SOP manual are protected works under the Copyright Act, 1957 and the Information Technology Act, 2000.')

# =====================================================================
# 13. (merged into 14) -- skip
# 14. DATA PROTECTION & CONFIDENTIALITY
# =====================================================================
H1('13. DATA PROTECTION AND CONFIDENTIALITY')
CL('13.1', 'Roles: The Company is the Data Fiduciary and alone determines the purposes and means of processing Customer Data. The Franchisee processes Customer Data only as the Company’s Data Processor, on the Company’s documented instructions, solely to operate the Store, and acquires no proprietary interest in it. The Data Processing Addendum at Annexure E governs such processing. The goodwill in customer relationships, and all anonymised business data, belong to the Company.')
CL('13.2', 'The Franchisee shall not export, share, sell, retain or use Customer Data for its own or any third party’s purpose, and shall comply with the Digital Personal Data Protection Act, 2023 and the Information Technology Act, 2000 as applicable.')
CL('13.3', 'Confidentiality: The Franchisee shall keep all Confidential Information secret during and after the Term, shall use it only to perform this Agreement, and shall procure the same from its staff. This Clause and Annexure E survive termination.')
CL('13.4', 'Data minimisation and misuse: The Company may provide Customer Data to the Franchisee only in masked or tokenised form through the Haper System (for example, call-masking and tokenised delivery addresses); the Franchisee shall not be given, and shall not attempt to obtain or export, bulk or raw customer personal data. Any unauthorised retention, export or use of Customer Data is both a breach of the Digital Personal Data Protection Act, 2023 and prohibited solicitation under Clause 14.2, and the Franchisee shall pay liquidated damages (a genuine pre-estimate) of ₹50,000 per proven instance, without prejudice to any other remedy.')

# =====================================================================
# 15 -> 14. NON-COMPETE & NON-SOLICITATION
# =====================================================================
H1('14. NON-COMPETE AND NON-SOLICITATION')
CL('14.1', 'During the Term, the Franchisee shall not, directly or indirectly, operate, manage, own (other than up to 1% of a listed company) or hold any financial interest in any competing quick-commerce or grocery-delivery business within the Territory.')
CL('14.2', 'The Franchisee acknowledges that the Customer Data, vendor data, SOPs and operational know-how are the Company’s exclusive Confidential Information. At all times during and after the Term the Franchisee shall not use or exploit such Confidential Information for any competing business; and for 12 (twelve) months after termination shall not, directly or indirectly, (a) solicit, divert or service any customer, vendor, employee or delivery person whose identity became known to it through such Confidential Information, or (b) solicit or employ any employee or delivery person engaged by the Company.')
CL('14.3', 'The Parties agree that Clause 14.2 protects the Company’s proprietary information, Customer Data and goodwill and is not a restraint of trade. Each limb of this Clause is severable and independent.')

# =====================================================================
# 16 -> 15. INDEMNITY & INSURANCE
# =====================================================================
H1('15. INDEMNITY AND INSURANCE')
CL('15.1', 'The Franchisee shall indemnify and hold the Company harmless against all claims, losses, penalties, fines and reasonable legal costs arising from the Franchisee’s breach, negligence, fraud, violation of law, handling of inventory/cash/data, or its employees’ or delivery persons’ acts or disputes.')
CL('15.2', 'The Franchisee shall maintain, at its cost, adequate (a) store/asset and stock insurance (naming the Company as loss-payee for Company-owned inventory and assets) for a sum not less than the value of the Company-owned stock and assets at the Store, (b) employee insurance (Workmen’s Compensation/ESI), and (c) public-liability insurance of at least ₹10,00,000. The Franchisee shall produce certificates of cover before go-live and on each renewal; if any cover lapses, the Company may take out such insurance and recover the premium from the Franchisee.')

# =====================================================================
# 17 -> 16. LIMITATION OF LIABILITY
# =====================================================================
H1('16. LIMITATION OF LIABILITY')
CL('16.1', 'Save for fraud, or for death or personal injury caused by its negligence, the Company shall not be liable for any indirect, special, incidental, consequential or punitive loss, or for loss of profit, revenue, goodwill, business opportunity or data.')
CL('16.2', 'The Company’s aggregate liability under or in connection with this Agreement shall not exceed the total net Sales Commission and Support Amount actually retained by the Company in the 12 (twelve) months preceding the event giving rise to the claim, and shall exclude any inventory sale proceeds, COD remittances, the Security Deposit and any amounts held in trust.')

# =====================================================================
# 18 -> 17. FORCE MAJEURE
# =====================================================================
H1('17. FORCE MAJEURE')
CL('17.1', '"Force Majeure Event" means any event beyond a Party’s reasonable control, including act of God, fire, flood, epidemic or pandemic, lockdown, government order, war, terrorism, riot, strike (other than of the affected Party’s own staff), and failure of utilities, internet, telecommunications, payment infrastructure or suppliers due to the foregoing.')
CL('17.2', 'Neither Party is liable for any failure or delay (other than in payment obligations) to the extent caused by a Force Majeure Event, provided it gives prompt notice and mitigates. A Force Majeure Event does not entitle the Franchisee to any early refund of the Security Deposit, to terminate during the lock-in period, or to relief from accrued sums. If it continues beyond 120 (one hundred and twenty) days, either Party may terminate without liability save for accrued obligations and the surviving clauses.')

# =====================================================================
# 19 -> 18 ... renumber termination as 18
# =====================================================================
H1('18. TERMINATION')
CL('18.1', 'Immediate Termination for Cause: The Company may terminate this Agreement immediately, by written notice and notwithstanding the lock-in period, for fraud, theft, brand misuse, inventory or sales manipulation, diversion of cash, criminal activity, insolvency of the Franchisee, or loss of any licence under Clause 8.4 / 9. These grounds require no cure period.')
CL('18.2', 'Events of Default (curable): Each of the following is an "Event of Default": non-payment of any sum due for more than 7 (seven) days; breach of Clauses 6, 8, 9, 11.4, 12, 13 or 14; failure to keep the Store operational, stocked and licensed; or any other material breach. On an Event of Default the Company may serve a written default notice; the Franchisee shall cure within 15 (fifteen) days (3 (three) business days for monetary default). If the breach is not cured (or is incapable of cure), the Company may terminate immediately by written notice, notwithstanding the lock-in period.')
CL('18.3', 'Termination for Convenience: After the lock-in period, either Party may terminate for convenience by giving 3 (three) months’ (90 days’) prior written notice; a notice by the Franchisee may not take effect during the lock-in period. The Security Deposit is refunded after termination in accordance with Clause 19.')
CL('18.4', 'Franchisee’s Right for Company Breach: The Franchisee may terminate for the Company’s uncured material breach (including failure to provide the platform for a continuous period exceeding 15 (fifteen) days, or failure to pay undisputed payouts within 30 (thirty) days of the due date) on 30 (thirty) days’ written notice if the Company fails to cure within that period.')
CL('18.5', 'Operational Closure: The Company shall give at least 30 (thirty) days’ notice before any final operational closure, where applicable.')
CL('18.6', 'Termination is without prejudice to accrued rights, the indemnities, the surviving clauses and the Personal Guarantee.')
CL('18.7', 'Step-in: If the Franchisee abandons the Store or ceases operations, or the Store is suspended (including on loss or suspension of any licence), or an Event of Default occurs, the Company may, without terminating, enter and take operational control of the Store, secure its inventory and assets, and operate or appoint another operator pending cure or termination, at the Franchisee’s cost. Such step-in is not a termination and is without prejudice to the Company’s other rights.')

# =====================================================================
# 20. REFUND & EARLY-EXIT
# =====================================================================
H1('19. REFUND OF SECURITY DEPOSIT ON TERMINATION')
CL('19.1', 'Refund: On termination of this Agreement by either Party, the Company shall refund the Security Deposit to the Franchisee, after deducting (with an itemised statement): (a) the depreciation in the value of the Company-funded store infrastructure and assets attributable to the Franchisee’s tenure, calculated per the depreciation schedule in Annexure F; (b) the documented landed cost of damaged, missing or expired stock attributable to the Franchisee; and (c) any other outstanding dues, finalised service-level charges / liquidated sums and permitted Deductions owed by the Franchisee.')
CL('19.2', 'GST on refund: As no GST is charged on the Security Deposit, no credit note or GST adjustment arises on its refund. Recovery of the cost of damaged, missing or expired stock under Clause 19.1(b) is the Franchisee’s indemnity for the Company’s loss. The infra-depreciation deduction under Clause 19.1(a) is treated as a charge for the use of Company assets, on which the Company shall raise a tax invoice and account for GST at the applicable rate.')
CL('19.3', 'Forfeiture / recovery on the Franchisee’s default: Where the Company terminates on any ground under Clause 18.1 (fraud, theft, inventory or sales manipulation, diversion of cash, criminal activity, insolvency, or loss of licence) or for an uncured material breach under Clause 18.2, the Company may apply the whole Security Deposit towards its losses by way of damages; and where its losses exceed the Security Deposit, recover the balance from the Franchisee (and the Guarantor, if any).')
CL('19.4', 'Settlement timeline: The Company shall complete the Full and Final (FnF) settlement — the joint stock reconciliation, finalisation of deductions, and return of all Company property and Customer Data — within 30 (thirty) days of the effective date of termination, and shall refund the net Security Deposit to the Franchisee within a maximum of 3 (three) months after the FnF settlement. The deductions in Clause 19.1 are without prejudice to the Company’s right to recover any shortfall exceeding the Security Deposit from the Franchisee (and the Guarantor, if any).')
CL('19.5', 'Support Amount clawback on early closure: The Support Amount (Annexure A) is a retention incentive that the Franchisee earns by operating the Store for the full Initial Term. If the Franchisee terminates this Agreement or closes the Store, or the Company terminates for the Franchisee’s default, before the end of the Initial Term, the Franchisee shall repay to the Company a portion of the aggregate Support Amount received by it, as follows: (a) if closure occurs before the completion of 24 (twenty-four) months from the Effective Date — 100% (one hundred per cent) of the aggregate Support Amount received; (b) if closure occurs on or after 24 months but before 36 (thirty-six) months — 50% (fifty per cent) of the aggregate Support Amount received; and (c) if closure occurs on or after the completion of 36 months — Nil. The Parties agree that this clawback is the repayment of an unearned incentive and a genuine pre-estimate of the Company’s loss, and not a penalty. The clawback amount is a due recoverable as a permitted Deduction and from the Security Deposit, and, to the extent it exceeds the Security Deposit, from the Franchisee (and the Guarantor, if any) under Clauses 4.5 and 19.4. This Clause does not apply where the Company terminates for convenience or for its own reasons not attributable to the Franchisee’s default. The clawback reflects the Company’s unrecovered sunk costs (including the rent and setup it funds for the full Term) and is a genuine pre-estimate of loss; where the original Support bore any GST or reverse charge, the clawback shall be given effect by a credit note under Section 34 of the CGST Act, 2017, and no fresh supply or tax invoice by the Franchisee arises on it.')
CL('19.6', 'No double recovery: The aggregate the Company recovers under Clauses 19.1, 19.3 and 19.5 and through the security cheques, the e-NACH mandate and any Personal Guarantee shall not exceed the Company’s actual loss together with its genuine pre-estimated loss; the Company shall not recover the same loss more than once, and any amount recovered under one head reduces the others pro tanto.')

# =====================================================================
# 21 -> 20. POST-TERMINATION
# =====================================================================
H1('20. POST-TERMINATION OBLIGATIONS')
CL('20.1', 'On termination the Franchisee shall immediately cease all use of the Licensed Marks and the Haper System and, within 7 (seven) days, de-identify the Store, remove all signage and branding, transfer or relinquish any Haper-related domain, app, marketplace, Google or social-media listing/handle, and refrain from any colourable imitation of the brand.')
CL('20.2', 'The Franchisee shall return all unsold inventory and Company assets in saleable/good condition, and shall return and then securely delete all Customer Data and Confidential Information (and certify deletion) within 7 (seven) days, as set out in Annexure E.')
CL('20.3', 'The Franchisee acknowledges that breach of this Clause causes irreparable harm and that the Company is entitled to injunctive relief.')

# =====================================================================
# 22 -> 21. ASSIGNMENT / CHANGE OF CONTROL
# =====================================================================
H1('21. ASSIGNMENT AND CHANGE OF CONTROL')
CL('21.1', 'The Franchisee shall not assign, novate, sub-franchise, sub-licence, charge or transfer this Agreement or any right or obligation under it, nor part with possession or operation of the Store, without the Company’s prior written consent.')
CL('21.2', 'Any change (in one or a series of transactions) in the legal or beneficial ownership of, or management control over, the Franchisee — including any transfer exceeding 26% (twenty-six per cent) of the voting/partnership interest, or any change in the controlling shareholder/managing partner — is deemed an assignment requiring such consent, which the Company may condition on a fresh Personal Guarantee and revised commercial terms. The Company shall have a right of first refusal exercisable within 30 (thirty) days.')
CL('21.3', 'The Company may assign this Agreement freely to any affiliate or successor, and may reorganise or novate it to a compliant affiliate for regulatory, tax or FEMA reasons provided the Franchisee’s commercial entitlements are not materially diminished, to which the Franchisee consents in advance.')
CL('21.4', '"Permitted assigns" means only an assignee approved in writing under this Clause who has executed a deed of adherence and a fresh Personal Guarantee. Any transfer in breach is void and a material breach.')

# =====================================================================
# 23 -> 22. REPRESENTATIONS & WARRANTIES
# =====================================================================
H1('22. REPRESENTATIONS, WARRANTIES AND COMPLIANCE')
CL('22.1', 'The Franchisee represents and warrants that it has full authority to enter into this Agreement; that its execution does not conflict with any other obligation; that the information and KYC documents it has provided are true; that its funds are from lawful sources and not in breach of the Prevention of Money-Laundering Act, 2002; and that neither it nor its principals is disqualified or debarred under law.')
CL('22.2', 'Anti-Bribery: Each Party shall comply with the Prevention of Corruption Act, 1988 and all applicable anti-bribery laws, and shall not offer or accept any improper payment in connection with this Agreement.')
CL('22.3', 'The Franchisee shall comply with all Applicable Law in operating the Store; the Company may suspend the Store on the platform pending cure of any material non-compliance.')

# =====================================================================
# 24 -> 23. RELATIONSHIP + FOREIGN INVESTMENT RESERVED MATTER
# =====================================================================
H1('23. RELATIONSHIP OF THE PARTIES')
CL('23.1', 'The Franchisee is an independent contractor and the Company’s commission agent for the limited purpose of Clause 6 only. Nothing in this Agreement creates a partnership, joint venture, or employer-employee relationship between the Company and the Franchisee or its staff. Neither Party may bind the other except as expressly provided.')

# =====================================================================
# 25 -> 24. DISPUTE RESOLUTION
# =====================================================================
H1('24. GOVERNING LAW AND DISPUTE RESOLUTION')
CL('24.1', 'Governing Law: This Agreement is governed by the laws of India.')
CL('24.2', 'Negotiation and Mediation: The Parties shall first attempt to resolve any dispute by good-faith negotiation and, failing that, by mediation — a time-boxed, directory step of 30 (thirty) days that does not bar arbitration or interim relief.')
CL('24.3', 'Arbitration: Any unresolved dispute shall be finally resolved by a sole arbitrator appointed by mutual written agreement of the Parties within 30 (thirty) days of a notice of arbitration, failing which by the High Court of Karnataka under Section 11 of the Arbitration and Conciliation Act, 1996. Neither Party may unilaterally appoint the arbitrator. The seat is Bengaluru; the language is English; and for claims up to ₹50,00,000 the arbitration shall be conducted on a fast-track, documents-only basis under Section 29B.')
CL('24.4', 'Court Relief (carve-out): Notwithstanding the agreement to arbitrate, (a) either Party may seek interim or conservatory measures from a competent court under Section 9 of the said Act; (b) the Company may institute proceedings (including a summary suit under Order XXXVII CPC) for injunctive or specific relief to protect its IP, brand, Confidential Information or Customer Data, for recovery or return of inventory and assets, and for recovery of admitted dues and enforcement of the Personal Guarantee; and (c) where inventory or assets are physically located outside Bengaluru (including at the Store in Bihar or elsewhere), the Company may additionally approach the court having jurisdiction over that location for their urgent protection or recovery. The seat of arbitration remains Bengaluru.')
CL('24.5', 'Jurisdiction: Subject to the above, the courts at Bengaluru have exclusive jurisdiction.')

# =====================================================================
# 26 -> 25. NOTICES
# =====================================================================
H1('25. NOTICES')
CL('25.1', 'Notices shall be in writing and sent to the addresses on the cover page / Annexure B (and to email addresses notified by the Parties) by hand, registered post, courier or email. A notice is deemed delivered on actual receipt, or 48 (forty-eight) hours after dispatch by registered post/courier, or on the day of sending by email (if no bounce is received), whichever is earlier.')

# =====================================================================
# 27 -> 26. MISCELLANEOUS
# =====================================================================
H1('26. MISCELLANEOUS')
CL('26.1', 'Entire Agreement: This Agreement and its Annexures constitute the entire agreement and supersede all prior proposals and understandings.')
CL('26.2', 'Amendments: Any amendment must be in writing and signed by both Parties. The Company may, however, update the SOPs, brand standards and platform pricing/discounts from time to time on reasonable notice.')
CL('26.3', 'Severability and Blue-Pencil: If any provision is held invalid or unenforceable, it shall be read down or severed to the minimum extent necessary, and the remainder continues in full force.')
CL('26.4', 'No Waiver: No failure or delay in exercising any right is a waiver of it.')
CL('26.5', 'Set-off and Interest: The Company may set off any amount due to it against any amount payable to the Franchisee. Overdue amounts carry interest at 18% (eighteen per cent) per annum until paid.')
CL('26.6', 'Survival: Clauses 4.5, 6.1, 6.4 to 6.7, 12, 13, 14, 15, 16, 17, 19, 20, 22, 23, 24, 25 and 26, and Annexures C, E and F, survive termination.')
CL('26.7', 'Counterparts and Electronic Execution: This Agreement may be executed in counterparts and by electronic signature, which are valid under Sections 5 and 10A of the Information Technology Act, 2000.')
CL('26.8', 'Personal Guarantee: The Company’s obligations are conditional on delivery of the duly executed and stamped Deed of Personal Guarantee at Annexure C, together with evidence of the guarantor’s authority.')

# =====================================================================
# EXECUTION
# =====================================================================
H1('IN WITNESS WHEREOF')
P('the Parties have executed this Agreement on the day and year first written above.')
P('')
P('For HAPVERSE PRIVATE LIMITED', bold=True)
P('Authorised Signatory: ____________________________')
P('Name: ____________________________    Designation: ____________________________')
P('')
P('For the FRANCHISEE', bold=True)
P('Signatory / Proprietor: ____________________________')
P('Name: ____________________________')
P('')
P('WITNESSES:', bold=True)
P('1. Name & Signature: ____________________________    Address: ____________________________')
P('2. Name & Signature: ____________________________    Address: ____________________________')

doc.add_page_break()

# =====================================================================
# ANNEXURE A — REVENUE TABLE
# =====================================================================
TITLE('ANNEXURE A — REVENUE SHARING TABLE', 13)
P('Monthly payout structure based on GMV achieved. "Commission Rate" is the commission as a percentage of GMV. Total Payout = Sales Commission + Support Amount − permitted Deductions (Clause 11.5).')

rows = [
 ('Monthly Sales (GMV)','Commission Rate (% of GMV)','Sales Commission','Support Amount','Total Payout'),
 ('₹3,00,000','5.00%','₹15,000','₹12,000','₹27,000'),
 ('₹3,50,000','4.88%','₹17,080','₹12,000','₹29,080'),
 ('₹4,00,000','4.75%','₹19,000','₹11,000','₹30,000'),
 ('₹4,50,000','4.62%','₹20,790','₹11,000','₹31,790'),
 ('₹5,00,000','4.50%','₹22,500','₹10,000','₹32,500'),
 ('₹5,50,000','4.38%','₹24,090','₹10,000','₹34,090'),
 ('₹6,00,000','4.25%','₹25,500','₹9,000','₹34,500'),
 ('₹6,50,000','4.12%','₹26,780','₹9,000','₹35,780'),
 ('₹7,00,000','4.00%','₹28,000','₹9,000','₹37,000'),
 ('₹7,50,000','3.88%','₹29,100','₹9,000','₹38,100'),
 ('₹8,00,000','3.75%','₹30,000','₹9,000','₹39,000'),
 ('₹8,50,000','3.62%','₹30,770','₹9,000','₹39,770'),
 ('₹9,00,000','3.50%','₹31,500','₹9,000','₹40,500'),
 ('₹9,50,000','3.38%','₹32,110','₹9,000','₹41,110'),
 ('₹10,00,000','3.25%','₹32,500','₹9,000','₹41,500'),
 ('₹10,50,000','3.12%','₹32,760','₹9,000','₹41,760'),
 ('₹11,00,000','3.00%','₹33,000','₹9,000','₹42,000'),
 ('₹15,00,000','3.00%','₹45,000','₹9,000','₹54,000'),
 ('₹20,00,000','3.00%','₹60,000','₹9,000','₹69,000'),
 ('₹30,00,000','3.00%','₹90,000','₹9,000','₹99,000'),
]
table = doc.add_table(rows=len(rows), cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = val
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
P('')
P('Note: For GMV above ₹30,00,000, the Commission Rate remains 3% with ₹9,000 fixed Support Amount. The Support Amount is a contractual incentive that may be reduced or suspended only as provided in Clause 11.3, and is a retention incentive subject to clawback on early closure under Clause 19.5 (100% if the Store closes before 24 months; 50% if between 24 and 36 months; nil thereafter). The Company bears the store rent directly to the landlord under Clause 10.3; the Franchisee bears no rent and the Support Amount is not a rent reimbursement.', italic=True, size=9)

doc.add_page_break()

# =====================================================================
# ANNEXURE B — TERRITORY
# =====================================================================
TITLE('ANNEXURE B — STORE AND TERRITORY SCHEDULE', 13)
P('Store Name (trading): ____________________________')
P('Store Address: ____________________________________________________________________')
P('Assigned Territory: a radius of 2 (two) kilometres from the Store Address. [The Company may set any radius between 2 and 3 kilometres; strike out and initial if different: ____ km.]')
P('Licensed Marks — registration particulars (to be inserted): ')
BULLET('"HAPER" — Trade Mark Application/Registration No.: ____________  Class(es): ____________  Status: ____________')
BULLET('"HAPVERSE" — Trade Mark Application/Registration No.: ____________  Class(es): ____________  Status: ____________')

doc.add_page_break()

# =====================================================================
# ANNEXURE C — PERSONAL GUARANTEE
# =====================================================================
TITLE('ANNEXURE C — DEED OF PERSONAL GUARANTEE', 13)
P('(To be executed on its own stamp paper of value prescribed under the applicable State Stamp Act, dated, and attested by two witnesses.)', italic=True, size=9)
P('IMPORTANT — use this guarantee ONLY where a THIRD-PARTY guarantor is available (e.g., a creditworthy family member). Where the Franchisee is an individual / sole proprietor, a guarantee from that same person adds nothing (he/she is already personally liable under Clause 4.5). If no third-party guarantor is available, this Annexure may be left unused, and the Company relies on the Security Deposit and the security cheques / e-NACH under Clause 4.5.', italic=True, size=9)
P('This DEED OF GUARANTEE is executed on ____________________ by the person(s) named in Schedule G-1 (the "Guarantor", and if more than one, jointly and severally) in favour of HAPVERSE PRIVATE LIMITED (the "Company"), in connection with the Franchise Agreement dated ____________________ (the "Agreement") between the Company and ____________________ (the "Franchisee").')
CL('1.', 'Guarantee: The Guarantor unconditionally and irrevocably guarantees the due performance of all obligations and the payment of all dues of the Franchisee under the Agreement. The Guarantor’s liability is joint and several with the Franchisee and co-extensive under Section 128 of the Indian Contract Act, 1872, and the Guarantor is additionally liable as primary obligor and principal debtor, such that this Guarantee remains enforceable even if any obligation of the Franchisee is or becomes void, voidable or unenforceable.')
CL('2.', 'No Exhaustion: The Company may proceed against the Guarantor without first proceeding against the Franchisee or enforcing any other security.')
CL('3.', 'Continuing Guarantee: This is a continuing guarantee under Section 129 covering all present and future dues for the Term and any renewal. Notwithstanding Sections 130 and 131, it is irrevocable, is not determined by notice, death or incapacity, and binds the Guarantor’s heirs, executors, administrators and estate.')
CL('4.', 'Consent to Variation: The Guarantor consents in advance to, and shall not be discharged by, renewal on revised commercial terms (Cl 3.2), additional expansion investment (Cl 3.3), changes to commission, Support, deductions, penalties or SOPs, any indulgence, time or composition granted to, or release of, the Franchisee, or the loss or non-perfection of any other security; and waives Sections 133, 134, 135, 139 and 141 of the Indian Contract Act, 1872 to the fullest extent permitted by law.')
CL('5.', 'Demand and Evidence: The Guarantor’s liability arises on the Company’s written demand (from which limitation runs). A statement of dues issued by the Company is conclusive evidence save for manifest error.')
CL('6.', 'Solvency and Survival: The Guarantor represents that it is solvent (Schedule G-2) and shall not dissipate its assets while any liability subsists. This Guarantee survives any insolvency, dissolution or resolution of the Franchisee.')
CL('7.', 'Subrogation: The Guarantor shall not exercise any right of subrogation until the Company has been paid in full.')
CL('8.', 'Dispute Resolution: Any dispute under this Deed shall be finally resolved by arbitration (sole arbitrator, seat Bengaluru, language English) on the same terms as Clause 24 of the Agreement, and Clauses 24 and 25 of the Agreement are incorporated into and form part of this Deed. The Guarantor expressly submits to, and by signing this Deed below, executes that arbitration agreement and the exclusive jurisdiction of the courts at Bengaluru.')
P('')
P('SCHEDULE G-1 — Guarantor(s): Name: ____________  S/o,D/o,W/o: ____________  Address: ____________  PAN: ____________  DIN (if any): ____________', size=9)
P('SCHEDULE G-2 — Solvency / net-worth declaration: ____________________________', size=9)
P('')
P('Guarantor Signature: ____________________   Name: ____________________', bold=True)
P('Witness 1: ____________________     Witness 2: ____________________')

doc.add_page_break()

# =====================================================================
# ANNEXURE D — SOP
# =====================================================================
TITLE('ANNEXURE D — OPERATIONAL STANDARDS (SOP SUMMARY)', 13)
P('The Franchisee shall ensure:')
BULLET('Inventory Handling: products stored in temperature-controlled conditions as required; FIFO strictly followed; near-expiry stock quarantined and returned/exchanged per SOP.')
BULLET('Customer Service: orders packed within 3 minutes of receipt; polite, professional conduct by all staff and delivery personnel.')
BULLET('Delivery Standards: at least 95% of orders delivered within the applicable service window, subject to lawful and safe riding (delivery timelines do not authorise any traffic or safety violation).')
BULLET('Store Hygiene: daily cleaning logs maintained; staff in Haper-branded uniforms; grooming and food-safety hygiene standards observed.')
BULLET('Reporting: daily stock and cash reconciliation uploaded to the Admin Dashboard before end-of-day.')
BULLET('Compliance: all municipal, fire-safety, food-safety and labour regulations complied with at all times.')
P('Service-Level Charges (genuine pre-estimates, not penalties): the Company may publish in the SOPs reasonable, pre-agreed charges for defined service-level failures (e.g., late/short reconciliation, repeated SLA breaches), which are recoverable as permitted Deductions under Clause 11.5.', italic=True, size=9)

doc.add_page_break()

# =====================================================================
# ANNEXURE E — DPA
# =====================================================================
TITLE('ANNEXURE E — DATA PROCESSING ADDENDUM', 13)
P('This Addendum governs the Franchisee’s processing of Customer Data as the Company’s Data Processor under Clause 13.')
CL('1.', 'Instructions: The Franchisee shall process Customer Data only on the Company’s documented instructions and solely to operate the Store. Any use for its own or a third party’s purpose is a material breach.')
CL('2.', 'Security: The Franchisee shall implement reasonable security safeguards — encryption in storage and transit, role-based access, activity logs retained for at least 1 (one) year, and an incident-response process.')
CL('3.', 'Breach Notification: The Franchisee shall notify the Company of any actual or suspected personal-data breach within 6 (six) hours of becoming aware, with all information the Company needs to meet its obligations. The Company alone notifies the Data Protection Board, data principals and CERT-In.')
CL('4.', 'Sub-processing and Transfer: No sub-processor and no transfer of Customer Data outside India without the Company’s prior written consent and compliance with any restriction under Section 16 of the DPDP Act, 2023.')
CL('5.', 'Return and Deletion: On termination, the Franchisee shall cease processing, return all Customer Data, and securely delete every copy (including backups) within 7 (seven) days, and certify deletion. Retention is a continuing breach and prohibited solicitation.')
CL('6.', 'Staff: Every staff member with access shall sign a confidentiality undertaking (surviving employment). Wrongful disclosure may attract liability under Section 72A of the Information Technology Act, 2000.')
CL('7.', 'Indemnity: The Franchisee indemnifies the Company, to the extent of its or its staff’s causative breach, for penalties, fines, compensation and costs.')

doc.add_page_break()

# =====================================================================
# ANNEXURE F — INFRASTRUCTURE & DEPRECIATION SCHEDULE
# =====================================================================
TITLE('ANNEXURE F — INFRASTRUCTURE & DEPRECIATION SCHEDULE', 13)
P('(To be itemised and signed by both Parties at store setup. This is used to calculate the depreciation deduction from the Security Deposit under Clause 19.1(a). Company-funded items remain the Company’s property under Clause 10.)', italic=True, size=9)
fcols = ('Company-funded item','Cost (₹)','Useful life (yrs)','Method')
frows = [
 fcols,
 ('Refrigeration / cold storage','____________','5','Straight-line'),
 ('Racks & shelving','____________','7','Straight-line'),
 ('POS & electronic devices','____________','3','Straight-line'),
 ('Signage & branding','____________','3','Straight-line'),
 ('Fit-out / electrical / civil','____________','7','Straight-line'),
 ('Other: ____________________','____________','__','Straight-line'),
 ('TOTAL','____________','',''),
]
ft = doc.add_table(rows=len(frows), cols=4)
ft.style = 'Light Grid Accent 1'
for i, row in enumerate(frows):
    for j, val in enumerate(row):
        c = ft.cell(i, j); c.text = val
        pr = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run('')
        pr.font.size = Pt(9)
        if i == 0 or (i == len(frows)-1): pr.bold = True
P('')
P('Depreciation method: straight-line over the useful life shown above (or such rate as the Parties initial against each item). On termination, the deduction from the Security Deposit under Clause 19.1(a) equals the accumulated depreciation of the above assets from the date of installation to the date of termination. The assets remain the Company’s property and no additional charge is made for them beyond this depreciation.', size=9)
P('')
P('Agreed — For the Company: ____________________     For the Franchisee: ____________________     Date: ____________', bold=True)

doc.add_page_break()

# =====================================================================
# FILL CHECKLIST
# =====================================================================
TITLE('SCHEDULE OF ITEMS TO FILL / DECIDE BEFORE SIGNING', 13)
P('(Internal checklist for the Company — remove THIS page and the note on the first page before handing a signed copy to the franchisee. Company particulars (CIN/PAN/TAN/GSTIN/bank) are already filled in the body.)', italic=True, size=9)

checklist = [
 ('A. Fill in the blanks (cover page & party block)', [
   'Execution Date and Effective Date on the cover page. If the Effective Date is left blank, it defaults to the date the FIRST deposit instalment is received in cleared funds.',
   'Franchisee details in the party block (Clause 2): full name, father/spouse name, age, PAN, Aadhaar (print ONLY the last 4 digits), and residential address. GSTIN only if the franchisee is actually registered.',
 ]),
 ('B. Complete the Annexures', [
   'Annexure B: Store name & full address; Territory radius (default 2 km, settable up to 3 km); and the "Haper"/"Hapverse" trade-mark application/registration numbers once filed.',
   'Annexure C (Personal Guarantee): use ONLY if a THIRD-PARTY guarantor is available — a sole proprietor guaranteeing his own debt is pointless. If used, fill the guarantor’s name, PAN, address and solvency declaration; stamp, date and witness it separately. If no third party, leave it unused and rely on the deposit + security cheques.',
   'Annexure F (Infrastructure & Depreciation): list the actual infra items and costs at store setup, have BOTH parties sign it, and take dated photos at handover — this is what makes the depreciation deduction clear and enforceable.',
 ]),
 ('C. Collect from the franchisee (before go-live)', [
   'KYC: PAN, a recent photograph, an independent address proof (NOT Aadhaar-derived), a CANCELLED CHEQUE (needed to refund the deposit), and only a MASKED Aadhaar (last 4 digits) or other government photo-ID. Do NOT collect/store the full Aadhaar number (UIDAI 2025 + DPDP); offer PAN + voter/passport as an alternative.',
   '3 SIGNED security cheques (with the franchisee’s written authority to fill in date + amount up to the dues, presented with a statement of dues) AND a signed e-NACH mandate (Clause 4.5) — your recovery tool for dues above the deposit, plus s.138 / PSS-Act leverage. Avoid blank/undated-only cheques.',
 ]),
 ('D. Security Deposit — money handling', [
   '₹7,50,000 refundable, interest-free SECURITY DEPOSIT (NOT a fee) — collect in 3 equal instalments of ₹2,50,000 by bank transfer: on signing, +15 days, +30 days.',
   'Issue a RECEIPT (not a tax invoice) for each instalment; charge NO GST; record it as a refundable LIABILITY (not income).',
   'DO NOT stock inventory or go live until all 3 instalments are received in full.',
   'Keep a cash reserve — the deposit is refundable within ~3 months of FnF on exit, so do not deploy 100% of every franchisee’s deposit.',
 ]),
 ('E. Set up with your CA / accountant (one-time)', [
   'Bihar GST registration — your GSTIN is Karnataka (29…); bill Bihar customers under a Bihar GSTIN. Moving stock Karnataka→Bihar is a SUPPLY between distinct persons (Schedule I): raise a TAX INVOICE + IGST + e-way bill (NOT a delivery challan); the IGST is creditable in Bihar. Repeat registration for each new State.',
   'ISD registration is MANDATORY since 1 Apr 2025 to distribute common third-party input-service credit (Bangalore SaaS/marketing) to Bihar — file GSTR-6 monthly; use cross-charge only for internally-generated HO services. TDS on franchisee payouts = s.194H @2% (not s.194-O).',
   'Confirm the GST treatment of the two exit deductions: damaged-stock recovery = indemnity, but reverse the input credit on written-off stock (s.17(5)(h)); the infra-depreciation deduction may attract GST as a charge for use of assets.',
   'Store lease (Clause 10.3): take a REGISTERED 3-year lease in Hapverse’s name — pay the stamp duty + registration and the landlord’s deposit. If the landlord is unregistered, account for 18% GST on rent under REVERSE CHARGE (RCM rule since Oct-2024). The lease in Hapverse’s name also supports billing Bihar customers under your Bihar GSTIN (store = your place of business). Get the landlord’s NOC for commercial + food (FSSAI) use, Haper branding, and assignment within the network.',
 ]),
 ('F. Brand & funding (one-time, before the first franchisee)', [
   'Register "Haper" as a trade mark in the Company’s name — your MCA company name does NOT give brand rights — and add the numbers to Annexure B.',
   'Funding: Indian-funded only, so no FDI action is needed now. If you ever take foreign capital, restructure first (the inventory-ownership model bars FDI in B2C e-commerce).',
 ]),
 ('G. Stamp & sign', [
   'Stamp THIS Agreement AND, separately, the Deed of Personal Guarantee — under the Stamp Act of the State of execution (Karnataka if signed at Bengaluru; Bihar if signed in Bihar). The STORE LEASE is a separate Bihar-situated instrument: register + stamp it in Bihar (≈6% of average annual rent) before the local Sub-Registrar, regardless of where it is signed. Sign with 2 witnesses and date all.',
   'Before handing the signed copy to the franchisee, DELETE this checklist page and the internal note on the first page.',
 ]),
 ('H. Latest-law compliance — Company’s own duties (verify with counsel)', [
   'Gig-worker laws are LIVE in BOTH your states: the Bihar Platform-Based Gig Workers Act 2025 AND the Karnataka Gig Workers Act 2025 (welfare fee 1–5% per transaction from 13 Feb 2026). Hapverse likely qualifies as the "aggregator" — get a classification opinion, register, and budget the welfare fee for delivery riders, regardless of the franchise structure.',
   'FSSAI: as an inventory e-commerce food entity, Hapverse needs its OWN central e-commerce FBO licence + a storage/dark-store licence per premises, and must follow the 30%-shelf-life-on-delivery norm and the CP (E-Commerce) Rules seller disclosures (grievance officer, country-of-origin, MRP) — these are NOT discharged by pushing them onto the franchisee.',
   'DPDP Act 2023: substantive obligations bite ~May 2027 — treat full build-out (consent notices, breach response, retention schedule) as a dated readiness project; Annexure E already covers the processor contract.',
   'POSH Act 2013: Hapverse (HO Bengaluru) must constitute its own Internal Committee — a Company compliance item separate from the franchisee’s.',
 ]),
]
for header, pts in checklist:
    H2(header)
    for pt in pts:
        BULLET(pt)

out = '/Users/office/Documents/haper/Hapverse Franchise Agreement - REDRAFTED.docx'
doc.save(out)
print('SAVED:', out)
